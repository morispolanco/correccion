import requests
import docx
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def translate_text(text, source_lang, target_lang):
    url = "https://ai-translate.pro/api/9428f69325adc980cc9b9dc6a0f84a30a3eb86e74787792c581cc44e4c1adfae/{}/{}".format(source_lang, target_lang)
    headers = {
        "Content-Type": "application/json"
    }
    data = {
        "text": text
    }
    response = requests.post(url, headers=headers, json=data)

    if response.ok and response.status_code == 200:
        translation = response.json()
        return translation["text"]
    else:
        print("Error en la solicitud de traducción.")
        return None

def compare_documents(original_doc, translated_doc):
    original_text = ""
    translated_text = ""

    # Leer el contenido del documento original
    with Document(original_doc) as doc:
        for paragraph in doc.paragraphs:
            original_text += paragraph.text + "\n"

    # Leer el contenido del documento traducido
    with Document(translated_doc) as doc:
        for paragraph in doc.paragraphs:
            translated_text += paragraph.text + "\n"

    # Comparar los dos textos
    if original_text == translated_text:
        print("Los documentos son idénticos.")
    else:
        print("Los documentos son diferentes.")

        # Guardar el control de cambios en un nuevo documento
        diff_doc = Document()
        diff_doc.add_heading("Control de cambios", level=1)
        diff_doc.add_paragraph("Documento original:", style="Heading2")
        diff_doc.add_paragraph(original_text)
        diff_doc.add_paragraph("Documento traducido:", style="Heading2")
        diff_doc.add_paragraph(translated_text)

        # Guardar el documento con el control de cambios
        diff_doc.save("control_de_cambios.docx")
        print("Se ha guardado el control de cambios en el archivo 'control_de_cambios.docx'.")

def main():
    # Cargar el documento original
    original_doc = "documento_original.docx"

    # Traducir el documento al inglés
    translated_text = translate_text(original_doc, "es", "en")

    if translated_text is not None:
        # Guardar el texto traducido en un nuevo documento
        translated_doc = "documento_traducido.docx"
        doc = Document()
        doc.add_paragraph(translated_text)
        doc.save(translated_doc)

        # Volver a traducir el documento al español
        retranslated_text = translate_text(translated_text, "en", "es")

        if retranslated_text is not None:
            # Guardar el texto re-traducido en un nuevo documento
            retranslated_doc = "documento_retraducido.docx"
            doc = Document()
            doc.add_paragraph(retranslated_text)
            doc.save(retranslated_doc)

            # Comparar los documentos original y re-traducido
            compare_documents(original_doc, retranslated_doc)

if __name__ == "__main__":
    main()
